from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFilter, ImageFont


ROOT = Path("/home/huge/dev/DGH")
THESIS_DIR = ROOT / "论文"
ASSET_DIR = THESIS_DIR / "generated_assets"
OUTPUT_PATH = THESIS_DIR / "快修宝-校园社区设备报修与工单跟踪系统毕业设计说明书.docx"


def font_path(name: str) -> str:
    candidates = {
        "regular": [
            Path("/mnt/c/Windows/Fonts/msyh.ttc"),
            Path("/mnt/c/Windows/Fonts/simsun.ttc"),
            Path("/mnt/c/Windows/Fonts/simhei.ttf"),
        ],
        "bold": [
            Path("/mnt/c/Windows/Fonts/msyhbd.ttc"),
            Path("/mnt/c/Windows/Fonts/simhei.ttf"),
            Path("/mnt/c/Windows/Fonts/simsun.ttc"),
        ],
    }
    for path in candidates[name]:
        if path.exists():
            return str(path)
    raise FileNotFoundError("未找到可用中文字体，请确认 WSL 可访问 Windows 字体目录。")


FONT_REGULAR = font_path("regular")
FONT_BOLD = font_path("bold")


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def set_run_font(run, font_name="宋体", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def style_paragraph(paragraph, first_line_chars: float = 2.0):
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(21 * first_line_chars)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(26)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, False)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", 15, True)
        p.paragraph_format.space_before = Pt(26)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
    elif level == 2:
        set_run_font(run, "黑体", 14, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.right_indent = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
    else:
        set_run_font(run, "黑体", 12, True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.right_indent = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(26)
    return p


def add_center_paragraph(doc: Document, text: str, size=12, bold=False, font_name="宋体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_text_block(doc: Document, text: str):
    for paragraph in [part.strip() for part in text.strip().split("\n\n") if part.strip()]:
        add_body_paragraph(doc, paragraph)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "黑体", 10.5, False)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(26)


def set_cell_text(cell, text: str, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "宋体", size, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], widths: Iterable[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(hdr[index], header, bold=True, size=10.5)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.JUSTIFY if index > 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[index], value, bold=False, size=10.5, align=align)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    add_caption(doc, title)
    return table


def add_placeholder(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_text(cell, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(18)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "右键目录后选择“更新域”，即可刷新页码。"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_text)
    run._r.append(fld_char_end)


def add_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, "宋体", 10.5)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_text)
    run._r.append(fld_char3)
    run = paragraph.add_run(" 页")
    set_run_font(run, "宋体", 10.5)


def rounded_box(draw, box, fill, outline, radius=28, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(draw, start, end, color, width=6):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head_len = 18
    spread = math.pi / 7
    left = (
        end[0] - head_len * math.cos(angle - spread),
        end[1] - head_len * math.sin(angle - spread),
    )
    right = (
        end[0] - head_len * math.cos(angle + spread),
        end[1] - head_len * math.sin(angle + spread),
    )
    draw.polygon([end, left, right], fill=color)


def wrap_text(draw, text: str, font, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        test = current + char
        box = draw.textbbox((0, 0), test, font=font)
        if box[2] - box[0] <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_text_block(draw, box, title: str, lines: list[str], fill_title, fill_text):
    title_font = load_font(34, bold=True)
    text_font = load_font(24, bold=False)
    x1, y1, x2, y2 = box
    draw.text((x1 + 24, y1 + 20), title, font=title_font, fill=fill_title)
    current_y = y1 + 72
    for line in lines:
        wrapped = wrap_text(draw, line, text_font, x2 - x1 - 48)
        for part in wrapped:
            draw.text((x1 + 24, current_y), part, font=text_font, fill=fill_text)
            current_y += 38


def canvas(size=(1600, 900)):
    image = Image.new("RGB", size, "#fffaf4")
    for i in range(3):
        layer = Image.new("RGBA", size, (0, 0, 0, 0))
        ld = ImageDraw.Draw(layer)
        if i == 0:
            ld.ellipse((60, 40, 560, 420), fill=(247, 234, 220, 180))
        elif i == 1:
            ld.ellipse((1050, 60, 1540, 430), fill=(255, 240, 216, 170))
        else:
            ld.ellipse((520, 520, 1320, 1060), fill=(233, 247, 255, 180))
        image = Image.alpha_composite(image.convert("RGBA"), layer).convert("RGB")
    return image.filter(ImageFilter.GaussianBlur(0.2))


def generate_tech_route(path: Path):
    image = canvas()
    draw = ImageDraw.Draw(image)
    accent = "#0f766e"
    accent_dark = "#0b4f4a"
    warn = "#f59e0b"
    rose = "#be123c"
    rounded_box(draw, (70, 180, 370, 730), "#ffffff", "#d8cfc2", radius=30)
    draw_text_block(draw, (70, 180, 370, 730), "业务角色", ["普通用户：提交报修、查看进度、评价反馈", "维修人员：签到、处理、完成工单", "管理员：派单调度、统计分析、预警处理"], accent_dark, "#2f2418")
    rounded_box(draw, (470, 110, 1110, 360), "#ffffff", "#d8cfc2", radius=30)
    draw_text_block(draw, (470, 110, 1110, 360), "前端表现层", ["Vue 3 + TypeScript + Vant + Vite", "移动端优先界面、三端路由、主题切换、表单交互", "用户端、维修端、管理端统一通过 Axios 访问接口"], accent, "#2f2418")
    rounded_box(draw, (470, 420, 1110, 720), "#f3fffb", "#85c8be", radius=30)
    draw_text_block(draw, (470, 420, 1110, 720), "后端服务层", ["NestJS + TypeORM + Class Validator", "控制器接收请求，服务层处理工单、签到、评价、统计逻辑", "定时扫描超时工单，保留自动派单和扩展指标接口"], accent, "#173129")
    rounded_box(draw, (1200, 140, 1510, 350), "#fff7ea", "#efc874", radius=30)
    draw_text_block(draw, (1200, 140, 1510, 350), "数据层", ["当前可运行版本：SQLite", "迁移预案：MySQL", "缓存预留：Redis"], warn, "#4a3620")
    rounded_box(draw, (1200, 420, 1510, 700), "#fff2f6", "#f0a7bd", radius=30)
    draw_text_block(draw, (1200, 420, 1510, 700), "运维支撑", ["Docker Compose 启停脚本", "建表与初始化数据脚本", "WSL 演示环境与前端构建产物"], rose, "#4a2430")
    draw_arrow(draw, (370, 455), (470, 235), accent)
    draw_arrow(draw, (370, 455), (470, 560), accent)
    draw_arrow(draw, (1110, 235), (1200, 245), "#43877f")
    draw_arrow(draw, (1110, 570), (1200, 560), "#43877f")
    image.save(path)


def generate_role_usecase(path: Path):
    image = canvas()
    draw = ImageDraw.Draw(image)
    title_font = load_font(30, bold=True)
    text_font = load_font(22)
    columns = [("普通用户", ["故障上报", "附带图片与位置", "查看工单进度", "提交维修评价"]), ("维修人员", ["查看待处理工单", "开始处理", "现场签到", "更新完成状态"]), ("管理员", ["查看统计看板", "人工派单", "处理超时预警", "分析高频故障点"])]
    colors = ["#0f766e", "#0ea5e9", "#be123c"]
    for idx, (name, items) in enumerate(columns):
        x1 = 90 + idx * 500
        x2 = x1 + 410
        rounded_box(draw, (x1, 180, x2, 740), "#ffffff", "#d8cfc2", radius=32)
        draw.text((x1 + 30, 210), name, font=title_font, fill=colors[idx])
        y = 290
        for item in items:
            draw.ellipse((x1 + 30, y + 10, x1 + 48, y + 28), fill=colors[idx])
            for line in wrap_text(draw, item, text_font, 300):
                draw.text((x1 + 70, y), line, font=text_font, fill="#2f2418")
                y += 34
            y += 18
    rounded_box(draw, (420, 60, 1180, 130), "#f3fffb", "#85c8be", radius=24)
    draw.text((470, 80), "三类角色围绕同一张工单协同，形成“上报—处理—复盘”闭环", font=load_font(26, True), fill="#0b4f4a")
    draw_arrow(draw, (500, 460), (590, 460), "#0ea5e9")
    draw_arrow(draw, (1000, 460), (1090, 460), "#be123c")
    image.save(path)


def generate_architecture(path: Path):
    image = canvas()
    draw = ImageDraw.Draw(image)
    rounded_box(draw, (140, 80, 1460, 240), "#ffffff", "#d8cfc2", radius=34)
    draw_text_block(draw, (140, 80, 1460, 240), "表示层", ["登录注册页、首页看板、报修表单、工单列表、维修工作台、管理看板、个人中心"], "#0f766e", "#2f2418")
    rounded_box(draw, (220, 330, 1380, 540), "#f3fffb", "#85c8be", radius=34)
    draw_text_block(draw, (220, 330, 1380, 540), "业务层", ["路由与权限判断、工单服务、签到服务、反馈服务、统计服务、超时扫描、自动派单预留逻辑"], "#0f766e", "#173129")
    rounded_box(draw, (300, 640, 1300, 820), "#fff7ea", "#efc874", radius=34)
    draw_text_block(draw, (300, 640, 1300, 820), "数据层", ["Ticket 与 CheckIn 实体持久化、初始化数据脚本、位置与用户扩展表设计、统计汇总数据来源"], "#d97706", "#4a3620")
    draw_arrow(draw, (800, 240), (800, 330), "#0f766e")
    draw_arrow(draw, (800, 540), (800, 640), "#d97706")
    image.save(path)


def generate_lifecycle(path: Path):
    image = canvas()
    draw = ImageDraw.Draw(image)
    steps = [("故障上报", "#0ea5e9"), ("待处理", "#f59e0b"), ("派单/自动分配", "#0f766e"), ("处理中", "#14b8a6"), ("现场签到", "#8b5cf6"), ("已完成", "#16a34a"), ("评价与统计", "#be123c")]
    x = 85
    y = 350
    for idx, (name, color) in enumerate(steps):
        box = (x + idx * 210, y, x + idx * 210 + 170, y + 120)
        rounded_box(draw, box, "#ffffff", color, radius=28, width=5)
        draw.text((box[0] + 22, box[1] + 40), name, font=load_font(24, True), fill=color)
        if idx < len(steps) - 1:
            draw_arrow(draw, (box[2], y + 60), (box[2] + 38, y + 60), color)
    rounded_box(draw, (180, 100, 1420, 240), "#f3fffb", "#85c8be", radius=30)
    draw_text_block(draw, (180, 100, 1420, 240), "状态控制要点", ["工单创建后进入待处理；派单后至少转为处理中；维修签到记录现场轨迹；完成后开放评分并进入统计口径。"], "#0b4f4a", "#173129")
    rounded_box(draw, (510, 640, 1090, 810), "#fff2f6", "#f0a7bd", radius=30)
    draw_text_block(draw, (510, 640, 1090, 810), "超时提醒", ["服务端按固定周期扫描超过 24 小时未完成工单，并输出预警信息，为后续短信或企业微信接入提供支点。"], "#be123c", "#4a2430")
    image.save(path)


def generate_deployment(path: Path):
    image = canvas()
    draw = ImageDraw.Draw(image)
    rounded_box(draw, (90, 280, 350, 610), "#ffffff", "#d8cfc2", radius=30)
    draw_text_block(draw, (90, 280, 350, 610), "访问终端", ["移动浏览器", "PC 浏览器", "答辩演示环境"], "#0f766e", "#2f2418")
    rounded_box(draw, (470, 190, 810, 420), "#ffffff", "#d8cfc2", radius=30)
    draw_text_block(draw, (470, 190, 810, 420), "前端服务", ["Vite 开发模式", "Nginx 承载静态产物", "Axios 统一访问 /api"], "#0ea5e9", "#1d3350")
    rounded_box(draw, (470, 500, 810, 770), "#f3fffb", "#85c8be", radius=30)
    draw_text_block(draw, (470, 500, 810, 770), "后端服务", ["NestJS 应用", "TypeORM 实体映射", "定时任务与控制器接口"], "#0f766e", "#173129")
    rounded_box(draw, (960, 150, 1510, 370), "#fff7ea", "#efc874", radius=30)
    draw_text_block(draw, (960, 150, 1510, 370), "当前数据方案", ["SQLite 单文件数据库", "适合课程设计展示", "部署轻量，初始化成本低"], "#d97706", "#4a3620")
    rounded_box(draw, (960, 450, 1510, 760), "#fff2f6", "#f0a7bd", radius=30)
    draw_text_block(draw, (960, 450, 1510, 760), "扩展部署方案", ["Docker Compose 编排", "MySQL + Redis 组合", "保留 schema.sql、data.sql 与容器脚本"], "#be123c", "#4a2430")
    draw_arrow(draw, (350, 445), (470, 305), "#0ea5e9")
    draw_arrow(draw, (350, 445), (470, 630), "#0f766e")
    draw_arrow(draw, (810, 305), (960, 255), "#d97706")
    draw_arrow(draw, (810, 630), (960, 600), "#be123c")
    image.save(path)


def generate_theme_concept(path: Path):
    image = canvas()
    draw = ImageDraw.Draw(image)
    rounded_box(draw, (120, 120, 720, 780), "#fffaf4", "#d8cfc2", radius=36)
    draw.rectangle((160, 210, 680, 300), fill="#0f766e")
    draw.text((195, 235), "白天模式", font=load_font(30, True), fill="white")
    for i in range(3):
        rounded_box(draw, (180, 360 + i * 110, 650, 450 + i * 110), "#ffffff", "#e7dccf", radius=24)
    draw.text((210, 390), "暖色背景、卡片化信息层级、低压视觉节奏", font=load_font(24), fill="#2f2418")
    draw.text((210, 500), "适合用户快速填写报修表单与浏览工单进度", font=load_font(24), fill="#2f2418")
    draw.text((210, 610), "通过色彩与留白降低移动端阅读疲劳", font=load_font(24), fill="#2f2418")
    rounded_box(draw, (880, 120, 1480, 780), "#091221", "#334a76", radius=36)
    draw.rectangle((920, 210, 1440, 300), fill="#172a4f")
    draw.text((955, 235), "夜晚模式", font=load_font(30, True), fill="#dfe8ff")
    for i in range(3):
        rounded_box(draw, (940, 360 + i * 110, 1410, 450 + i * 110), "#13203a", "#27406e", radius=24)
    draw.text((970, 390), "深色底搭配青蓝与洋红高亮，形成夜间提示感", font=load_font(24), fill="#dfe8ff")
    draw.text((970, 500), "适合夜间值班与维修处理，减少强光刺激", font=load_font(24), fill="#dfe8ff")
    draw.text((970, 610), "把工单状态、告警和操作按钮做出更强区分度", font=load_font(24), fill="#dfe8ff")
    image.save(path)


def generate_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    generate_tech_route(ASSET_DIR / "01_技术路线.png")
    generate_role_usecase(ASSET_DIR / "02_角色用例.png")
    generate_architecture(ASSET_DIR / "03_总体架构.png")
    generate_lifecycle(ASSET_DIR / "04_工单生命周期.png")
    generate_deployment(ASSET_DIR / "05_部署拓扑.png")
    generate_theme_concept(ASSET_DIR / "06_双主题概念.png")


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.left_indent = Pt(0)
    normal.paragraph_format.right_indent = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(42)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(26)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_cover(doc: Document):
    for _ in range(4):
        doc.add_paragraph()
    add_center_paragraph(doc, "毕业设计说明书", size=22, bold=True, font_name="黑体")
    doc.add_paragraph()
    add_center_paragraph(doc, "题目：", size=14, bold=True, font_name="黑体")
    add_center_paragraph(doc, "“快修宝”——校园/社区设备报修与工单跟踪系统的设计与实现", size=18, bold=True, font_name="黑体")
    doc.add_paragraph()
    doc.add_paragraph()
    labels = ["学院：__________________", "年级：__________________", "专业：__________________", "学生姓名：__________________", "学号：__________________", "指导教师：__________________", "职称：__________________", "日期：__________________"]
    for label in labels:
        p = add_center_paragraph(doc, label, size=14, bold=False, font_name="宋体")
        p.paragraph_format.space_after = Pt(8)
    doc.add_page_break()


def add_abstracts(doc: Document):
    add_center_paragraph(doc, "摘  要", size=16, bold=True, font_name="黑体")
    add_text_block(doc, """
随着高校后勤管理和社区公共服务逐步向数字化、移动化方向演进，设备报修业务已经从单纯的信息登记问题，转变为兼顾响应时效、过程透明、责任追踪与数据复盘的综合管理问题。传统电话报修、纸面登记和口头转达方式存在信息缺失、派单效率低、维修过程不可见、评价数据难沉淀等痛点，既增加了后勤岗位的沟通成本，也削弱了用户对服务质量的感知。围绕这一现实场景，课题设计并实现了“快修宝”校园/社区设备报修与工单跟踪系统，仓库内部以 DGH 作为工程代号，通过统一的工单模型把报修上报、维修处置、管理调度和统计复盘连接为一条完整业务链。

系统采用移动端优先的前后端分离架构。前端基于 Vue 3、TypeScript、Vant 与 Vite 构建，围绕普通用户、维修人员和管理员三类角色设计差异化操作界面；后端基于 NestJS、TypeORM 与类验证机制实现工单、签到、评价和统计等核心业务，并通过定时扫描方式实现超时工单预警。为了兼顾答辩演示环境的轻量化和后续工程扩展的连续性，当前可运行版本采用 SQLite 保存核心业务数据，同时在工程中保留了 MySQL、Redis、Docker Compose 与初始化脚本等扩展部署方案，为系统后续迁移到更完整的生产化结构提供支撑。

论文围绕课题的需求分析、系统总体设计、详细设计、代码实现、测试验证和总结展开展示。实践结果表明，系统已经完成报修提交、工单查询、状态更新、维修签到、评价反馈、基础统计与超时提醒等主链路功能，能够支撑校园或社区设备报修场景下的核心业务闭环；同时，前端已预留登录注册、人员管理、洞察看板和双主题视觉等扩展页面，为进一步完善认证能力、用户治理、数据分析与运维看板打下界面基础。项目在 WSL 环境中完成前后端构建验证，说明其具备较好的可部署性和可维护性，能够满足毕业设计阶段对功能完整性、结构合理性和可演示性的要求。
""")
    p = doc.add_paragraph()
    style_paragraph(p, first_line_chars=0)
    run = p.add_run("关键词：设备报修；工单跟踪；Vue 3；NestJS；移动端优先；运维闭环")
    set_run_font(run, "宋体", 12, False)
    doc.add_page_break()


def add_sample_like_first_page(doc: Document):
    for _ in range(8):
        doc.add_paragraph()
    add_center_paragraph(doc, "毕业设计说明书", size=54, bold=True, font_name="黑体")
    for _ in range(9):
        doc.add_paragraph()
    title_para = add_center_paragraph(doc, "“快修宝”——校园/社区设备报修与工单跟踪系统的设计与实现", size=18, bold=False, font_name="黑体")
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title_para.paragraph_format.line_spacing = Pt(26)
    info_para = add_center_paragraph(doc, "2022710318    邓甘豪    22级软件工程三班", size=12, bold=False, font_name="宋体")
    info_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    info_para.paragraph_format.line_spacing = Pt(26)
    teacher_para = add_center_paragraph(doc, "指导教师    陈睿", size=12, bold=False, font_name="宋体")
    teacher_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    teacher_para.paragraph_format.line_spacing = Pt(26)
    doc.add_page_break()


def add_toc_page(doc: Document):
    add_center_paragraph(doc, "目  录", size=16, bold=True, font_name="黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_toc(p)
    doc.add_page_break()


def add_figure(doc: Document, image_name: str, caption: str, width_cm: float = 15.5):
    doc.add_picture(str(ASSET_DIR / image_name), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def build_body(doc: Document):
    add_heading(doc, "一、 前言", 1)
    add_heading(doc, "（一）研究背景", 2)
    add_text_block(doc, """
校园和社区是典型的高频设备运维场景。照明、供水、空调、门禁、公共健身器材、教学设备等资产分布广、使用密集、故障类型杂，一旦管理方式仍停留在电话登记或人工转交阶段，就会产生信息记录碎片化、责任边界模糊、处理过程难追踪等连锁问题。尤其在人员有限、区域分散的后勤组织中，报修信息一旦描述不完整，派单人就需要反复回访确认，维修人员到场后又可能因为位置不清、问题不明而重复沟通，导致工单从生成到完结的全过程都被低效率所拖累。

设备报修系统的建设价值，不仅在于提供一个“线上表单”替代纸面登记，更在于通过工单机制重构服务过程。故障上报阶段要让信息尽量结构化，便于快速研判；派单阶段要让责任人清晰可见，降低任务漂移；维修阶段要保留签到和状态更新轨迹，形成过程证据；完成阶段要收集反馈和统计数据，为后续分析高频故障点、优化备件配置和改进服务策略提供基础。也就是说，设备报修系统本质上是后勤服务组织方式的信息化重塑，而不是简单的页面开发任务。

在智慧校园和智慧社区建设持续推进的背景下，移动端优先已经成为此类系统的重要特征。普通用户往往在故障现场发起报修，维修人员常常在巡检或处理途中更新状态，管理员则需要通过简洁看板掌握进展和风险。因此，一个适合毕业设计场景的设备报修系统，应当同时满足三项要求：其一，主链路功能清晰可运行；其二，界面操作适应手机端使用习惯；其三，工程结构保留继续扩展到更完整生产架构的可能。课题围绕这三点展开设计与实现。
""")
    add_heading(doc, "（二）国内外研究现状", 2)
    add_text_block(doc, """
从国外研究与应用情况看，设施管理系统、IT 服务管理系统和工单平台起步较早，很多产品已经形成成熟的事件上报、优先级划分、派单协同、服务等级管理和绩效分析机制。这类系统普遍强调流程标准化，通过工单状态机、通知机制和知识库沉淀来提高运维效率。在高校和大型物业场景中，设备报修往往与资产台账、巡检计划、供应商协同等模块联动，形成更完整的运维平台。其优势在于流程完善、数据维度丰富，但同时也存在部署成本高、定制门槛高的问题，对于中小规模教学项目或轻量级演示场景并不总是合适。

国内相关研究更多聚焦于“Spring Boot + Vue”“移动端优先”“校园后勤数字化”“工单可视化管理”等方向。近年高校管理系统、社区服务平台和设备维护平台的设计论文普遍强调前后端分离、角色权限控制、响应式界面和数据库结构优化等内容。随着前端技术和云部署方式的成熟，许多研究不再停留在后台管理系统层面，而是开始关注用户侧的交互体验、处理过程的透明度以及运维数据的可视化表达。这一趋势说明，设备报修系统已经从单纯的信息采集工具转向服务体验与管理效率并重的综合平台。

但从课程设计与毕业设计角度看，现有方案仍存在两个常见问题。第一，部分项目在论文中描述的功能远多于实际代码落地内容，导致设计与实现脱节；第二，部分项目虽然功能罗列较全，却忽略了业务闭环，只完成了“增删改查”而没有将上报、派单、处理、反馈和统计串联起来。课题在吸收既有研究经验的基础上，选择把工单主链路作为建设重点，先完成闭环，再为认证、人员管理、看板指标和缓存机制预留扩展位点，以保证说明书内容能够与项目仓库中的实际实现相互对应。
""")
    add_heading(doc, "（三）研究目标与主要内容", 2)
    add_text_block(doc, """
课题的直接目标，是设计并实现一套面向校园或社区场景的设备报修与工单跟踪系统，使普通用户能够方便地提交故障信息，维修人员能够按工单处理任务并留下过程记录，管理员能够掌握工单总量、完成情况、评价水平和超时风险。从软件工程视角看，这一目标还包含结构合理、模块清晰、便于测试和便于扩展四个要求，因此系统并未以单页展示为终点，而是通过前后端分离和实体建模组织项目代码。

围绕上述目标，论文的主要研究内容分为六个层面：一是梳理设备报修业务中的角色关系、流程节点和痛点，形成需求边界；二是结合仓库中现有代码、初始化脚本和部署材料，确定项目的技术路线与工程形态；三是完成系统总体设计，包括架构分层、模块划分、业务流程和数据库概念设计；四是展开详细设计，说明前端页面、后端服务、状态流转、接口结构与数据模型之间的对应关系；五是结合构建验证和核心流程测试，对系统实现效果进行分析；六是在总结阶段指出当前版本的价值、局限与后续扩展方向。

需要说明的是，课题推进过程中保留了两个层次的工程方案。其一，是已能够直接运行的核心版，即 Vue 3 + NestJS + SQLite 所支撑的工单闭环；其二，是结合 `schema.sql`、`data.sql`、`docker-compose.yml` 和前端预留页面所呈现的增强版规划，包括 MySQL/Redis 迁移、认证体系完善和人员治理能力扩展。论文对这两个层次会做出明确区分：凡属于当前代码已实现的功能，将在实现与测试章节按实际情况说明；凡属于界面与脚本中已经预留但尚未完全联调的能力，则归入扩展设计或后续规划，避免“纸面功能大于代码功能”的问题。
""")
    add_heading(doc, "（四）论文结构安排", 2)
    add_text_block(doc, """
全文共分为八个部分。第一部分为前言，说明课题背景、研究现状、研究目标与章节安排；第二部分介绍项目采用的关键技术、开发环境与工程路线；第三部分完成可行性与需求分析，明确角色、功能和非功能需求；第四部分给出系统总体设计，集中说明架构分层、模块划分、业务流程和数据概念模型；第五部分展开详细设计，对各核心模块与接口进行逐项说明；第六部分结合仓库代码分析系统实现路径和工程落地情况；第七部分展示测试环境、测试方法、主要用例和结果分析；第八部分对全文进行总结，并给出进一步改进方向。
""")

    add_heading(doc, "二、 相关技术与开发环境", 1)
    add_heading(doc, "（一）项目技术路线与架构选型", 2)
    add_text_block(doc, """
课题采用前后端分离技术路线。前端负责页面展示、角色导航、表单交互和状态可视化，后端负责业务规则处理、实体持久化与统计汇总，数据层负责保存工单与签到信息。与单体页面相比，这种方案有两方面优势：一方面可以把界面逻辑和业务逻辑拆开，降低耦合度，使页面改版不会直接影响服务层；另一方面有利于按模块迭代开发，前端可以围绕页面交互并行推进，后端则围绕实体和接口完成实现与测试。

仓库中同时保留了两类工程痕迹：README、Docker Compose 和配置文件中仍有 Spring Boot + MySQL + Redis 的部署规划，而实际可运行后端已经演进为 NestJS + TypeORM + SQLite 的轻量实现。这种差异并不意味着项目失控，恰恰反映了课题在开发中作出的工程取舍。为了在毕业设计阶段优先保证主链路可运行、环境搭建足够稳定，项目选择将后端落到 TypeScript 技术栈，并用 SQLite 降低安装成本；与此同时，又保留了 MySQL/Redis 结构设计和容器化脚本，便于后续向更完整的后端形态迁移。论文因此把“当前实现方案”和“扩展部署方案”同时纳入分析，但两者的定位严格区分。
""")
    add_figure(doc, "01_技术路线.png", "图2-1 系统技术路线示意")
    add_heading(doc, "（二）前端技术", 2)
    add_text_block(doc, """
前端部分采用 Vue 3 组合式 API 作为核心开发框架。与传统选项式写法相比，组合式 API 更适合在毕业设计项目中组织中等复杂度页面逻辑，能够把数据、事件和副作用按关注点集中管理，例如首页同时维护统计数据与洞察数据，报修页面同时维护表单状态与弹窗选择器，这类逻辑在 `script setup` 结构下更清晰。配合 TypeScript，页面开发不仅可以获得更严格的接口字段提示，还能在调试阶段提前暴露状态枚举、函数参数和接口返回结构上的问题。

UI 组件层面，项目使用 Vant 作为移动端组件库。设备报修业务大量依赖表单、按钮、弹窗、选择器、标签和图片预览等基础交互，Vant 在手机端的适配较好，既能减少基础控件的重复开发，又能保持视觉风格统一。项目中的报修页、登录注册页、列表页和管理端表单都借助 Vant 完成，说明组件化方案确实提高了开发效率。页面样式则由项目自定义的 `styles.css` 进一步强化，形成了“白天温暖、夜晚霓虹”的双主题视觉语言，使系统在答辩展示中不至于落入普通管理系统的单调模板化界面。

构建工具采用 Vite。Vite 在开发阶段提供更快的启动速度与热更新体验，适合毕业设计周期内频繁调整页面。构建结果表明，前端项目能够在 WSL 环境中顺利完成生产构建，说明其依赖关系和模块组织是稳定的。前端还通过 Axios 封装了统一的 `api` 实例，并在请求拦截器中保留鉴权令牌挂载逻辑，为后续接入后端认证能力提供了清晰入口。
""")
    add_figure(doc, "06_双主题概念.png", "图2-2 系统双主题视觉概念图")
    add_heading(doc, "（三）后端技术", 2)
    add_text_block(doc, """
后端采用 NestJS 构建。NestJS 在结构上延续了模块、控制器、服务等分层思想，非常适合课程设计中强调“模块化、可维护、便于扩展”的场景。项目中 `TicketsModule` 负责组织工单相关实体、控制器和服务，`TicketsController` 暴露 REST 风格接口，`TicketsService` 聚合创建工单、修改状态、派单、签到、评价、统计和超时查询等业务逻辑。这种结构让“接口入口”和“业务处理”保持清晰分离，既方便阅读，也便于后续针对某个模块增加验证或拆分子服务。

数据访问层使用 TypeORM。相较于手写 SQL 拼接，TypeORM 在毕业设计场景下具有两点实用价值。第一，它允许通过实体类直接描述字段结构、主键策略和关联关系，便于在论文中将“数据库设计”与“代码落地”对应起来；第二，Repository 与 QueryBuilder 的组合能够满足当前项目的 CRUD 操作和超时筛选需求，减少了样板代码。项目中的 `Ticket` 实体包含工单标题、描述、图片列表、位置、状态、报修人、指派人、评分与反馈等字段，`CheckIn` 实体则保留维修签到轨迹，两者之间通过一对多关系形成工单处理链条。

参数校验方面，项目使用 `class-validator` 与 `ValidationPipe` 形成请求数据过滤机制。创建工单、派单、签到、评价和状态修改都通过 DTO 约束字段类型和必填条件，有助于提升接口健壮性。应用启动时还通过 `setInterval` 周期性调用超时工单查询逻辑，从而实现轻量级定时提醒。这些实现虽然规模不大，却覆盖了一个后端服务的关键能力：输入校验、实体映射、业务编排和定时任务。
""")
    add_heading(doc, "（四）数据存储与扩展部署技术", 2)
    add_text_block(doc, """
当前可运行版本选用 SQLite 作为核心数据存储方案。其优点在于部署简单、文件级管理方便、演示环境搭建成本极低，不需要单独安装数据库服务即可完成工单主链路验证。对于毕业设计而言，这种选择有很强的现实意义：答辩现场和新机器上恢复运行环境的成本更低，也能减少因为外部服务配置不当而导致的演示风险。NestJS 通过 TypeORM 直连 `data.sqlite` 文件，即可完成数据表同步和基础持久化。

但 SQLite 并非课题最终扩展的唯一方向。仓库中的 `schema.sql` 与 `data.sql` 提供了面向 MySQL 的表结构和初始化数据，`docker-compose.yml` 则进一步给出 MySQL 与 Redis 的容器化编排方案。从数据库设计可以看出，项目已经预留了 `users`、`locations`、`tickets` 和 `check_ins` 四类表结构，除了核心工单与签到外，还支持位置维度和用户维度的数据组织；Redis 在 Compose 中作为缓存组件保留，虽当前版本尚未真正联调，但其存在表明项目在扩展层面已经考虑到高频状态读取、验证码缓存和通知队列等需求。

因此，论文对数据技术的表述采取“双层说明”策略：一方面，以 SQLite + TypeORM 作为当前实现的直接依据；另一方面，以 MySQL + Redis + Docker Compose 作为增强型部署路径，说明项目如何从课堂演示形态平滑过渡到更接近生产部署的形态。这样的表述既尊重仓库中的真实实现，也能体现系统在工程层面具有迁移和扩展的前瞻性。
""")
    add_heading(doc, "（五）开发环境与工具", 2)
    add_text_block(doc, """
项目开发与验证环境采用 Windows 主机配合 WSL 的方式完成。代码仓库存放在 WSL 文件系统中，前后端依赖安装与构建验证均在 Ubuntu 子系统内部执行，以避免 Windows 侧直接在 `\\\\wsl.localhost` 路径上运行 `npm` 所触发的 UNC 路径兼容问题。实际构建结果表明，前端在 WSL 中能够正常执行 `npm run build`，后端也能在完成依赖安装后顺利通过 TypeScript 编译，这一结果为论文中的部署与测试分析提供了直接依据。

具体工具方面，前端依赖 Node.js、npm、Vite 和浏览器完成调试，后端依赖 Node.js、npm、TypeScript 与 NestJS 运行时，文档生成则使用 Python、python-docx 与 Pillow 自动生成说明书和插图。版本管理可通过 Git 完成，脚本启动可依赖仓库自带的 `dev-up.sh`、`dev-status.sh` 和 `dev-down.sh`。从软件工程角度看，这套工具链覆盖了编码、构建、部署、文档与版本控制五个环节，足以支撑毕业设计的完整实施过程。
""")
    add_heading(doc, "三、 系统可行性与需求分析", 1)
    add_heading(doc, "（一）可行性分析", 2)
    add_heading(doc, "1. 技术可行性", 3)
    add_text_block(doc, """
从技术成熟度看，课题使用的 Vue 3、TypeScript、Vant、NestJS、TypeORM 与 SQLite 都属于文档充分、社区活跃、学习成本可控的主流技术。它们既能够支撑设备报修系统所需的页面交互、接口开发和数据持久化，也适合在毕业设计周期内完成从需求分析到代码落地的全流程开发。系统的业务复杂度主要集中在工单状态流转和多角色页面组织，而不是高并发事务或分布式一致性，因此现有技术方案足以胜任。

从实现难度看，项目将复杂度控制在“可被验证的闭环”范围内。后端只围绕 Ticket 与 CheckIn 两个核心实体建立关系，减少了过度建模带来的负担；前端则围绕三个角色工作台构建页面，以通用卡片、列表和表单承载业务操作，避免陷入繁重的组件封装工作。与此同时，项目又保留了认证拦截器、看板接口和容器化部署材料的接入点，使系统在主链路稳定后仍具备继续增长的空间。因此，从毕业设计的时间与能力边界衡量，该方案具有良好的技术可行性。
""")
    add_heading(doc, "2. 经济可行性", 3)
    add_text_block(doc, """
系统开发所依赖的软件框架和工具基本都可以免费获取，开发阶段不需要采购额外授权，也不依赖高成本云资源。答辩演示时采用 SQLite 作为数据库，可省去单独配置数据库主机或云容器的开支；若进入进一步扩展阶段，MySQL 与 Redis 也同样可以通过开源版本完成部署。对于毕业设计而言，这种成本结构具有明显优势，能够把主要精力集中在系统设计和实现本身，而不是投入到商业软件采购或复杂基础设施准备上。
""")
    add_heading(doc, "3. 运行与维护可行性", 3)
    add_text_block(doc, """
设备报修系统的使用群体以普通用户、维修人员和管理员为主，他们更关注流程是否顺畅、界面是否直观，而不会花太多时间学习复杂操作。因此系统必须具有较低的学习门槛。项目通过标签化状态展示、卡片式工单列表、移动端表单与按钮操作降低了使用阻力；通过统一接口路径和模块化服务实现降低了代码维护门槛。再加上仓库中保留了建表脚本、示例数据和容器脚本，说明该系统在运行维护层面也具有较强的可恢复性和可迁移性。
""")
    add_heading(doc, "（二）业务场景与角色分析", 2)
    add_text_block(doc, """
设备报修业务表面上只有“提交问题”和“维修处理”两个动作，实际上涉及三类利益主体。普通用户追求的是报修入口清晰、处理进度可知、结果反馈及时；维修人员关注的是任务来源明确、现场信息充分、状态更新方便；管理员关注的是工单总量、资源分配、超时风险和服务评价。若系统只服务其中一方，就无法形成闭环。例如只有用户报修页而没有维修签到和管理派单功能，最终仍需要依赖线下沟通；只有后台统计而没有前端表单规范化输入，统计数据也会失真。

因此，系统需求分析必须以角色协同为核心。用户端承担“问题发现与发起”职责，强调信息采集和结果感知；维修端承担“问题处置与过程留痕”职责，强调时效与轨迹记录；管理端承担“资源调度与复盘治理”职责，强调总览和干预。项目在页面与接口设计上都遵循这一分工原则，把多角色协作统一映射到同一张工单对象上。这样一来，系统中的每一次状态变化都能找到责任主体和业务含义，避免信息脱节。
""")
    add_figure(doc, "02_角色用例.png", "图3-1 三角色协同用例图")
    add_heading(doc, "（三）功能需求分析", 2)
    add_text_block(doc, """
普通用户端的核心需求包括故障上报、位置选择、附件补充、工单进度查看和服务评价。故障上报不是单纯录入一句描述，而是要包含标题、详细说明、位置、图片和联系方式等要素，尽量一次性把维修人员需要的信息说明清楚。工单查看需求则要求用户能看到状态变化、签到轨迹、指派人信息和评价入口，从而感知报修是否真正进入处理流程。评价功能用于收集服务质量反馈，为管理端统计提供数据来源。

维修端的核心需求包括工单浏览、状态更新、现场签到和处理反馈。维修人员通常处于移动作业环境中，系统应支持以最少步骤完成“开始处理”“完成工单”“签到记录”这类操作，因此页面必须避免复杂表单。签到不是附属功能，而是维修过程是否真实发生的重要凭据，它既能帮助管理员复核处理轨迹，也能为后续分析平均到场时间、区域工单负载等指标提供基础数据。

管理端的核心需求包括统计总览、工单派单、高频故障点识别和超时预警。管理员不直接维修，但需要知道工单是否积压、哪些区域故障频繁、哪些人员负载过高、哪些工单超过服务时限，因此系统必须提供基础统计能力。此外，前端还预留了人员管理、运营看板和洞察摘要等页面，这些页面代表的是产品化增强方向：当认证和用户表完全打通后，系统即可从单纯工单平台进一步演化为具有组织治理能力的后勤服务平台。
""")
    add_heading(doc, "（四）非功能需求分析", 2)
    add_text_block(doc, """
在性能方面，毕业设计阶段的设备报修系统不需要追求极端并发指标，但至少应保证在常规教学演示和小规模使用环境下拥有稳定响应。为此，系统采用结构较为轻量的接口设计，数据查询范围集中在工单与签到两类核心实体，能够满足一般场景下的流畅使用。若后续接入更多用户、位置和通知能力，则可以通过迁移至 MySQL、接入 Redis 缓存和增加索引等方式进一步提升性能。

在安全性方面，系统首先要保证输入数据的合法性，防止无效字段进入业务层；其次要具备明确的角色边界，避免用户越权访问他人数据；再次要保证数据库和接口之间的字段含义一致，避免“前端字段存在、后端字段缺失”造成的隐性风险。当前后端已经在创建工单、派单、签到、评价和状态更新接口上加入 DTO 校验，前端也在请求层预留了鉴权令牌挂载逻辑，为补全认证体系提供基础。

在可用性与可维护性方面，系统强调移动端优先、信息层级清晰、页面结构统一和模块边界明确。可用性体现在用户可在较少操作步骤内完成报修或处理动作；可维护性体现在前端按视图拆分、后端按模块组织、实体与 DTO 相互对应、样式集中管理。对于毕业设计项目而言，这种“结构先于功能堆砌”的思路，比表面上罗列大量模块更有实际价值，因为它决定了项目后续是否还能继续演化。
""")

    add_heading(doc, "四、 系统总体设计", 1)
    add_heading(doc, "（一）设计原则", 2)
    add_text_block(doc, """
系统总体设计遵循四项原则。第一，闭环优先原则。设备报修平台的价值来自完整流程，而非孤立页面，因此设计时优先打通“上报—派单—处理—签到—完成—评价—统计”这条主链路。第二，移动端优先原则。报修和维修动作大量发生在现场，操作界面必须适应手机端单手交互与碎片化使用场景。第三，轻量实现原则。在毕业设计周期和演示环境受限的前提下，优先选择安装成本低、联调路径短的技术方案，确保项目能够稳定落地。第四，可扩展原则。虽然当前实现聚焦核心链路，但系统结构中必须保留认证、用户管理、缓存、看板和容器化部署的扩展接口，为后续深化预留空间。
""")
    add_heading(doc, "（二）系统总体架构设计", 2)
    add_text_block(doc, """
系统总体架构可以概括为表示层、业务层和数据层三层结构。表示层由前端路由、页面组件与样式系统构成，负责接收用户输入、展示状态信息并组织多角色导航；业务层由 NestJS 控制器、服务、实体映射与校验机制构成，负责实现工单创建、派单、签到、评价、统计和超时查询；数据层则通过 SQLite 或未来可迁移的 MySQL 完成持久化存储，并为统计与扩展功能提供数据基础。三层之间仅通过明确接口进行交互，避免前端直接依赖数据库或业务层绕过实体模型。

这种架构的优势在于职责划分清楚。前端如果需要调整视觉布局，只要不改变接口契约，就不会影响后端；后端如果需要替换数据库类型或增加缓存，也不必大规模改动页面逻辑。对于毕业设计项目而言，明确的层次结构不仅有助于代码管理，也方便在说明书中进行逐层讲解，让系统设计能够被外部读者快速理解。
""")
    add_figure(doc, "03_总体架构.png", "图4-1 系统总体架构图")
    add_heading(doc, "（三）模块划分设计", 2)
    add_text_block(doc, """
从业务视角出发，系统可以拆分为用户报修模块、维修处理模块、管理调度模块、统计分析模块和系统支撑模块五部分。用户报修模块主要负责采集故障信息并发起工单；维修处理模块负责工单状态推进和签到留痕；管理调度模块负责指派责任人、查看总体工单情况并进行人工干预；统计分析模块负责汇总总工单数、已完成数量、平均评分和高频区域等指标；系统支撑模块则包含主题切换、路由权限、初始化脚本、超时扫描和部署脚本等基础能力。各模块之间以工单为共同核心对象，从而保证系统结构不会被割裂为互不关联的页面集合。

前端页面中的登录、注册、人员管理、运营概览和个人中心等内容，属于系统在产品层面的扩展模块。它们的意义在于把当前“核心工单系统”向“完整服务平台”方向延展：登录与注册为真实用户体系做准备，人员管理为角色治理做准备，运营概览为更深层指标分析做准备，个人中心则为通知偏好和账号维护做准备。虽然这些模块尚未全部与当前后端对齐，但在总体设计中将其纳入模块版图是合理的，因为它们对应的是系统自然演进方向，而不是与主题无关的附加功能。
""")
    add_heading(doc, "（四）业务流程总体设计", 2)
    add_text_block(doc, """
工单业务流程可以概括为七个节点。首先，普通用户在移动端填写故障标题、详细描述、位置、图片和联系方式，系统生成工单并赋予初始状态。其次，管理员或自动派单机制为工单指定维修人员，工单状态从待处理进入处理中。第三，维修人员查看工单详情并在到达现场后完成签到，系统记录时间和地点。第四，维修人员根据处置情况更新工单状态。第五，当故障解决后工单进入已完成状态。第六，用户在工单列表中对本次维修进行评分与文字反馈。第七，管理端汇总工单总量、完成量、评分和高频故障区域，并对超时工单进行预警处理。整个流程以状态变化为主线，以签到和评价为辅线，形成可追踪的服务链。
""")
    add_figure(doc, "04_工单生命周期.png", "图4-2 工单生命周期与提醒机制示意图")
    add_heading(doc, "（五）数据库概念设计", 2)
    add_text_block(doc, """
从概念模型看，系统最核心的实体是工单 Ticket。工单承担业务主轴，记录故障内容、报修位置、状态变化、报修人、指派人、评分和反馈等信息。围绕工单，一类是签到实体 CheckIn，用于记录维修人员到场和处理过程中的关键轨迹；另一类是用户与位置这类扩展实体，在 MySQL 预案中分别以 `users` 和 `locations` 表体现，用于支撑更加规范的账号体系和位置选择体系。当前 SQLite 版本在实体层面只实现了 Ticket 与 CheckIn，以降低演示复杂度；但从表结构规划看，项目已经具备向更完整数据模型扩展的基础。

数据库概念设计的重点，不是把实体数量做得很多，而是保证实体之间的关系能真实反映业务。工单与签到是一对多关系，因为一个工单可能对应多次现场记录；工单与用户在扩展方案中表现为“报修人”“维修人”和“管理员”围绕同一业务对象发生作用；工单与位置之间则体现为“故障发生在哪里”的空间维度。如果后续继续扩展，还可以增加设备台账实体，将“哪台设备发生故障”与“在哪个位置发生故障”区分开来，进一步提升系统的专业性。
""")
    add_placeholder(doc, "此处预留 ER 图，请在终稿中补充根据项目数据库设计绘制的 E-R 图。")
    add_caption(doc, "图4-3 数据库 E-R 图预留位置")
    add_heading(doc, "（六）接口与安全总体设计", 2)
    add_text_block(doc, """
系统接口遵循 REST 风格设计，路径围绕 `tickets` 和 `stats` 两类业务资源展开。这样设计的好处是资源边界清晰，前端路由和接口语义容易对应。例如新增工单使用 `POST /tickets`，查看工单列表使用 `GET /tickets`，状态修改通过 `PATCH /tickets/:id/status` 实现，签到与评价则作为工单的关联动作分别挂在 `checkins` 和 `feedback` 子路径下。这种接口结构能够直接反映业务语义，也方便在说明书中解释各接口的职责。

安全设计方面，当前可运行版本的重点在请求数据合法性和状态控制正确性；前端和扩展接口中则继续保留了登录态与令牌挂载逻辑。换言之，系统已经完成“输入安全”和“流程安全”的基础层建设，认证安全与用户治理将在扩展版本中进一步完善。这样的分阶段建设方式符合毕业设计项目的实际规律：先让核心业务正确运行，再逐步补强平台级能力，而不是一开始就引入过多复杂机制影响主链路落地。
""")
    add_heading(doc, "五、 系统详细设计", 1)
    add_heading(doc, "（一）前端路由、角色入口与主题设计", 2)
    add_text_block(doc, """
前端通过 Vue Router 组织多角色页面入口。未登录用户默认进入登录页或注册页，登录后根据角色自动跳转到用户首页、维修工作台或管理端主页。这样的设计减少了用户寻找入口的成本，也让系统天然具备角色隔离能力。虽然当前后端尚未完全打通认证接口，但前端已经在 `beforeEach` 守卫中实现了按角色跳转和拦截逻辑，这意味着项目在界面结构上已经考虑到真实使用场景，而不是简单地把所有页面平铺在一个菜单中。

主题设计是该项目前端的一项特色。系统通过 `theme.ts` 维护手动模式、自动模式和临时覆盖逻辑，在页面顶栏与个人中心中提供切换入口。白天主题采用暖纸色背景与柔和卡片层次，强调亲和感和可读性；夜晚主题则使用深色底加霓虹高亮，强调夜间值班场景下的对比度与警示性。这个设计并非纯粹的美化处理，而是把报修场景中“白天办公”和“夜间抢修”的不同使用氛围纳入了界面表达，有助于提升系统辨识度与答辩展示效果。
""")
    add_heading(doc, "（二）用户报修与工单查询模块设计", 2)
    add_text_block(doc, """
用户报修模块以 `ReportView.vue` 为核心页面，围绕“尽量一次性把问题说清楚”的目标组织表单。页面要求填写故障标题、详细描述、位置与报修人信息，并支持填写图片链接，前端还通过弹窗选择器承载位置和报修人选项。在产品化视角下，这一页面的价值不是简单收集几项字段，而是通过界面结构引导用户形成结构化描述。标题用于快速分类故障，描述用于补充细节，位置决定派单范围，图片有助于维修人员提前判断工具或备件需求。

工单查询模块则以 `TicketsView.vue` 为核心，提供全部、待处理、处理中、已完成和待评价等多种筛选条件。卡片中会展示工单标题、位置、时间、状态标签、报修人信息、维修员信息以及签到轨迹，用户在已完成但尚未评分的工单上还能直接提交评价。该设计将“查看结果”与“参与复盘”合并到一个界面中，减少页面跳转。尤其是时间轴式签到记录，使用户能够感知工单并非停留在系统中，而是已经被维修人员实际处理，这对于提升服务透明度具有直接作用。
""")
    add_heading(doc, "（三）维修处理与签到模块设计", 2)
    add_text_block(doc, """
维修工作台页面以待处理数量、处理中数量和今日完成数量作为顶部摘要，先让维修人员快速判断工作压力，再通过卡片列表展示各工单的故障内容和当前状态。界面提供“开始处理”“完成”“打卡签到”三个核心操作按钮，把维修过程浓缩为最常用的动作集合，避免现场人员被复杂表单拖慢节奏。与传统后台页面相比，这种设计更贴近移动维修人员的使用习惯。

签到功能在业务上具有双重价值。其一，它能记录维修人员是否已经到达现场，为工单处理过程提供证据；其二，它可以沉淀维修轨迹数据，为后续统计平均到场时长、区域响应速度和人员活跃度提供基础。在后端实现中，签到通过 `POST /tickets/:id/checkins` 接口写入 CheckIn 实体，与工单形成一对多关系。前端卡片中同步展示签到时间和地点，使管理员与报修人都能对维修进度形成更直观的判断。
""")
    add_heading(doc, "（四）管理调度与统计分析模块设计", 2)
    add_text_block(doc, """
管理端承担“总览”和“干预”两类职责。总览体现在统计面板、SLA 指标、高频区域和维修画像等板块中，干预则体现在对单个工单的人工派单能力。当前后端已经支持总工单数、已完成数、平均评分和高频位置统计，前端在此基础上进一步预留了运营闭环看板，包括超时工单数量、区域负责人建议、维修绩效画像和待回访评价数等内容。这种设计方式既保证了当前版本能够展示基础统计结果，也让系统在说明书中能够完整表达“数据驱动运维”的目标形态。

人工派单逻辑在页面上表现为管理员输入维修员姓名与联系方式后，对指定工单执行派单操作。后端服务会更新 `assignedToName` 和 `assignedToContact` 字段，并在工单处于待处理状态时自动推进为处理中。这一规则体现了系统对业务语义的理解：派单不是孤立的信息填充，而是流程状态的切换。通过把“指派责任人”和“进入处理中”绑定在一起，系统避免了工单已指定责任人但状态仍显示待处理的语义冲突。
""")
    add_heading(doc, "（五）工单状态控制与超时提醒设计", 2)
    add_text_block(doc, """
工单状态是整个系统最核心的控制字段。项目将其划分为 `pending`、`in_progress` 和 `completed` 三种枚举值，看似简单，实则覆盖了当前课程设计场景下最关键的业务阶段。创建工单时默认进入待处理；派单或自动指派后进入处理中；维修完成后进入已完成；评价功能仅在已完成状态下开放。这种设计保留了足够清晰的流程含义，又避免因为状态过多导致实现复杂度上升。

超时提醒则建立在状态控制之上。后端服务通过计算“当前时间减去工单创建时间是否超过 24 小时且状态不为已完成”来筛选超时工单，并在应用启动后按周期执行扫描。虽然当前实现采用日志提醒的轻量方式，但从工程结构看，这一机制已经具备继续扩展到短信、微信或企业 IM 通知的基础。也就是说，系统并没有把“超时工单”仅作为统计结果看待，而是把它作为主动触发运维干预的信号，这是设备报修系统走向实用化的重要一步。
""")
    add_heading(doc, "（六）数据库逻辑设计", 2)
    add_text_block(doc, """
在当前 SQLite 实现中，`Ticket` 表承担工单主体信息，字段涵盖标题、描述、图片列表、位置、状态、报修人、指派人、评分、反馈以及创建更新时间。图片字段使用文本存储 JSON 数组，这种处理方式简化了多图附件的保存逻辑，适合当前轻量化场景。`CheckIn` 表记录签到人、签到地点、关联工单和创建时间，通过外键关系与工单表连接，构成现场轨迹信息。

在 MySQL 扩展方案中，表结构进一步细化为 `users`、`locations`、`tickets` 和 `check_ins` 四类实体。这样做的意义在于把“字符串字段直接写入工单”的轻量方案，逐步过渡到“实体关联”的规范方案。例如位置不再是自由文本，而是由 `locations` 表统一维护；报修人也不再只是姓名和联系方式组合，而是由 `users` 表提供正式身份信息。数据库逻辑设计因此具有明显的演进路径：先以轻量实现保证主链路，再以规范实体支撑后续平台化扩展。
""")
    add_table(doc, "表5-1 核心工单表主要字段设计", ["字段名", "含义", "设计说明"], [["id", "工单主键", "采用 UUID，避免前后端并发创建时出现主键冲突，便于后续跨环境迁移。"], ["title", "故障标题", "用于列表快速识别问题类别，支持管理员进行人工判断和后续分类统计。"], ["description", "详细描述", "保存故障现象、影响范围与现场情况，是维修人员判断处置方式的重要依据。"], ["images", "图片数组", "当前以 JSON 文本形式保存，可在后续扩展到对象存储或上传文件表。"], ["location", "故障位置", "当前版本采用文本字段，扩展方案中可关联 locations 表实现标准化位置管理。"], ["status", "工单状态", "通过 pending、in_progress、completed 三值驱动整个流程的状态展示与操作权限。"], ["assignedToName / assignedToContact", "指派信息", "用于记录管理员指定的维修责任人，是工单追责与沟通联系的重要字段。"], ["rating / feedback", "服务评价", "工单完成后由用户填写，用于形成满意度统计和服务改进依据。"]], widths=[3.2, 3.8, 8.2])
    add_table(doc, "表5-2 维修签到表主要字段设计", ["字段名", "含义", "设计说明"], [["id", "签到记录主键", "采用 UUID，确保每一次现场签到都能被独立追踪。"], ["technicianName", "签到人姓名", "当前版本直接保存姓名，满足轻量化展示需求。"], ["location", "签到地点", "记录维修人员到达现场的实际位置，为过程核验提供凭据。"], ["ticket", "关联工单", "通过多对一关系连接 Ticket 实体，支持一个工单拥有多条签到轨迹。"], ["createdAt", "签到时间", "用于后续计算响应时长、过程时长和绩效指标。"]], widths=[3.2, 3.8, 8.2])
    add_heading(doc, "（七）接口设计", 2)
    add_text_block(doc, """
接口设计强调“操作动作围绕资源展开”，避免出现路径随意、语义含混的问题。核心资源为工单，因此新增、查询、详情、派单、状态更新、签到和评价都围绕 `tickets` 资源展开；统计则聚合到 `stats` 资源下。前端页面与接口之间基本呈一一映射关系：用户报修页对应工单创建接口，工单列表页对应工单查询接口，维修工作台对应状态更新与签到接口，管理端则对应派单与统计接口。这种映射关系有利于系统联调，也有利于论文说明。
""")
    add_table(doc, "表5-3 主要接口设计说明", ["接口", "方法", "作用"], [["/tickets", "POST", "创建新工单，接收标题、描述、图片、位置和报修人信息。"], ["/tickets", "GET", "获取工单列表，并连带返回签到记录。"], ["/tickets/{id}", "GET", "查询指定工单详情，用于查看完整状态与轨迹。"], ["/tickets/{id}/status", "PATCH", "更新工单状态，驱动待处理、处理中和已完成之间的切换。"], ["/tickets/{id}/assign", "PATCH", "为工单指派维修人员，并在必要时推进流程。"], ["/tickets/{id}/checkins", "POST", "写入现场签到记录，保留维修过程轨迹。"], ["/tickets/{id}/feedback", "POST", "提交评分与文字评价，形成服务反馈数据。"], ["/tickets/overdue/list", "GET", "查询超时未完成工单，为管理员预警与干预提供依据。"], ["/stats", "GET", "返回总工单、已完成数量、平均评分和高频位置等基础统计指标。"]], widths=[5.2, 2.2, 8.0])

    add_heading(doc, "六、 系统实现", 1)
    add_heading(doc, "（一）后端服务实现", 2)
    add_text_block(doc, """
后端代码组织相对集中，适合毕业设计说明书进行逐层映射。`main.ts` 完成应用启动、全局校验管道注册、CORS 开启和超时提醒定时器绑定；`app.module.ts` 负责加载 TypeORM 配置和工单模块；`ticket.entity.ts` 与 `checkin.entity.ts` 则分别定义业务实体；`tickets.controller.ts` 暴露接口；`tickets.service.ts` 承担实际业务逻辑。这样的结构让“入口、路由、规则、数据”四个层面清晰可辨，有利于维护和二次开发。

业务实现中最值得关注的是 `TicketsService`。创建工单时，服务会判断是否存在环境变量形式的自动派单配置，若存在则直接给新工单填入指派人并把状态推进到处理中；工单查询时会连带返回签到信息；派单时会根据当前状态自动推进流程；统计时会计算总量、完成量、平均评分和位置频次；超时筛选时则通过创建时间与状态组合条件进行查询。这说明当前后端虽规模不大，但已经形成了较完整的领域服务雏形，而不是简单地把每个接口都写成直接操作数据库的样板函数。
""")
    add_heading(doc, "（二）前端页面实现", 2)
    add_text_block(doc, """
前端页面实现以角色场景为组织单位。首页 `HomeView.vue` 重点展示统计摘要、热点区域、流程说明和三端协同价值；报修页 `ReportView.vue` 强调结构化表单填写；工单页 `TicketsView.vue` 强调状态筛选、列表浏览和评价提交；维修端 `TechnicianView.vue` 把摘要指标与现场操作整合在同一页面；管理端 `AdminView.vue` 则把统计卡片、闭环看板与工单派单能力统一组织。页面之间通过顶部导航和底部 TabBar 串联，形成符合移动端使用习惯的操作路径。

从实现效果看，项目在视觉与交互上表现出较强的一致性。统一的卡片圆角、标签颜色、状态徽标和主题色变量，使不同页面即使功能差异较大，也能保持同一产品语言。系统首页突出温暖服务和夜间守护的品牌意象，报修页强化信息完整性提示，维修端强调现场任务感，管理端强调数据全局感。这些设计细节说明前端并非简单堆砌组件，而是围绕报修场景对页面气质和信息层级进行了有针对性的构造。
""")
    add_placeholder(doc, "此处预留系统运行截图，可在终稿中补充首页、报修页、工单页、维修端和管理端的实际界面截图。")
    add_caption(doc, "图6-1 系统主要运行界面截图预留位置")
    add_heading(doc, "（三）部署与初始化实现", 2)
    add_text_block(doc, """
工程部署分为轻量运行方案和增强部署方案两类。轻量方案以 WSL 为基础，直接在前端和后端目录下安装依赖并执行开发或构建命令，适合毕业设计演示与日常调试；增强方案则借助 `docker-compose.yml` 同时编排 MySQL、Redis、后端和前端容器，适合后续扩展验证。两种方案并行存在的好处在于：轻量方案保证了当前项目可快速恢复运行，增强方案则保证了系统具备向更正式环境迁移的路径。

初始化方面，仓库中提供了 `schema.sql` 与 `data.sql`，可用于快速创建用户、位置、工单和签到等示例数据。即便当前 SQLite 版本未直接使用这套 MySQL 脚本，它们依然具有重要价值，因为这部分材料为论文中的数据库设计、角色示例与扩展部署提供了直接依据。对毕业设计项目而言，保留初始化脚本是一种较好的工程习惯，它意味着项目不依赖“开发者脑海中的隐性数据”，而是可以被外部环境重复恢复。
""")
    add_figure(doc, "05_部署拓扑.png", "图6-2 系统部署拓扑图")
    add_heading(doc, "（四）扩展功能与演进实现", 2)
    add_text_block(doc, """
项目目前已经表现出明显的产品演进思路。前端注册页、登录页、人员管理页、个人中心页以及洞察与运营看板接口，都说明系统并未把自己限定为一次性的 CRUD 练习，而是在现有工单闭环之上继续探索账号体系、用户治理、主题切换、通知偏好、绩效分析等更贴近真实平台的能力。这部分内容虽然尚未全部完成前后端闭环，但它们不是与主题无关的附会，而是围绕设备报修平台自然延伸出来的能力集合。

从工程管理角度看，把扩展需求先体现在前端路由、接口封装和配置材料中，再逐步补齐服务端，是一种符合实际开发节奏的做法。它能够让开发者尽早验证页面结构和产品逻辑，及时发现哪些模块值得优先落地，哪些模块应保留为下一阶段任务。论文将这种实现方式界定为“主链路先落地、扩展能力分阶段补齐”，这样既不会夸大当前版本的实际完成度，也能体现课题在工程规划上的完整性。
""")

    add_heading(doc, "七、 系统测试", 1)
    add_heading(doc, "（一）测试目标", 2)
    add_text_block(doc, """
系统测试的目标分为三个层次。第一层是功能正确性，验证工单创建、查询、状态更新、派单、签到、评价和统计是否与设计一致；第二层是流程完整性，验证工单是否能够从提交一路流转到评价和统计，而不是停留在中间节点；第三层是工程可运行性，验证前后端项目在当前环境下是否能够完成构建与部署。这三层目标对应毕业设计对“能实现、能运行、能说明”的核心要求。
""")
    add_heading(doc, "（二）测试环境与方法", 2)
    add_text_block(doc, """
测试环境采用 Windows 主机上的 WSL Ubuntu 子系统。前端在 WSL 中执行 `npm run build`，后端在 WSL 中执行 `npm install` 与 `npm run build`。之所以强调 WSL，是因为在 Windows 侧直接对 `\\\\wsl.localhost` 路径执行 npm 会遭遇 UNC 路径限制，导致命令默认切换到 Windows 目录并使依赖构建失败，这一现象在测试过程中已被实际观察到。将构建流程固定在 WSL 内部后，前端与后端都能顺利完成构建，因此论文在部署建议中明确推荐使用 WSL 方案。

测试方法以黑盒功能测试和工程构建验证为主。黑盒测试关注接口输入输出和页面操作结果是否符合预期；构建验证关注依赖安装、打包编译和运行环境是否稳定。考虑到毕业设计阶段的时间和资源边界，测试不以压力测试为核心，而是重点覆盖角色主流程、状态转换和常见异常场景，从而保证系统在展示与答辩中具备足够稳定性。
""")
    add_heading(doc, "（三）测试用例设计", 2)
    add_table(doc, "表7-1 主要测试用例设计", ["编号", "测试内容", "预期结果", "测试结论"], [["T1", "提交包含标题、描述、位置、联系人和图片数组的报修工单", "系统成功生成工单，状态默认为待处理或自动进入处理中", "通过"], ["T2", "查询工单列表并核对签到信息是否一并返回", "系统按创建时间倒序返回工单，已有关联签到数据可正常展示", "通过"], ["T3", "管理员对待处理工单执行派单", "系统成功写入维修员信息，并将工单推进到处理中", "通过"], ["T4", "维修人员对工单执行签到和状态更新", "系统生成签到记录，工单状态可切换为处理中或已完成", "通过"], ["T5", "已完成工单提交评分与反馈", "系统保存评分与文字评价，统计接口平均评分同步更新", "通过"], ["T6", "在 WSL 中构建前端项目", "Vite 成功输出 dist 目录，说明页面依赖与模块组织稳定", "通过"], ["T7", "在 WSL 中构建后端项目", "TypeScript 编译通过，说明服务端代码结构与依赖关系可正常工作", "通过"], ["T8", "在 Windows 侧直接对 WSL UNC 路径执行 npm 命令", "命令出现路径兼容问题，构建不稳定，不建议作为正式流程", "发现环境限制，已规避"]], widths=[2.0, 5.2, 6.0, 3.0])
    add_heading(doc, "（四）测试结果分析", 2)
    add_text_block(doc, """
从测试结果看，系统核心业务链已经具备较好的可用性。创建工单、查询列表、派单、签到、完成和评价这几类操作之间不存在明显断裂，说明工单模型的设计能够支撑业务闭环。统计接口能够基于现有工单数据计算总量、完成量、平均评分和高频位置，说明后端服务不只是保存数据，还具备基本的数据汇总能力。前端在 WSL 环境中的构建成功，则进一步证明其依赖关系和路由组织是稳定的。

测试也暴露出一个值得在说明书中明确记录的问题，即 Windows 侧直接对 `\\\\wsl.localhost` 路径运行 npm 命令时会因为 UNC 兼容限制而出现异常。这一问题并非业务代码错误，而是运行环境与命令执行方式之间的兼容性问题。通过将依赖安装和构建过程统一迁移到 WSL 内部，问题得到规避。这一经验说明，毕业设计项目不仅要关注业务功能本身，还应重视运行环境与工程流程的一致性，否则即便代码正确，也可能在演示阶段因环境问题影响结果。

综合来看，系统当前版本已经达到毕业设计要求中的“功能完整、结构清晰、能够运行与展示”的目标，但仍存在一些可继续提升的方向。例如认证能力仍需与后端完全对齐，用户与位置的规范化实体还需在当前 SQLite 版本中进一步落地，运营看板中的扩展接口也有待补齐。这些不足并不影响核心链路的完成度，反而表明系统仍具有继续深化和工程化完善的空间。
""")

    add_heading(doc, "八、 总结与展望", 1)
    add_text_block(doc, """
课题围绕校园或社区设备报修场景，完成了“快修宝”设备报修与工单跟踪系统的设计与实现。与单纯展示页面或孤立接口不同，系统始终以工单闭环为核心，将用户报修、维修处理、管理员派单、签到留痕、评价反馈和基础统计串联为统一流程。前端通过移动端优先设计和双主题风格增强了现场使用体验，后端通过模块化服务、实体映射和参数校验保障了业务逻辑的清晰性与稳定性。实际构建验证也证明项目能够在 WSL 环境中完成打包，具备较好的可运行性。

从毕业设计的完成度来看，本项目最大的价值在于做到了“设计与实现基本一致”。论文中涉及的核心功能都能在仓库代码中找到对应实现依据，而扩展能力也在前端页面、初始化脚本和部署配置中留下了明确痕迹，没有把尚未完成的内容包装成既成事实。这样的工程态度有助于说明书保持真实、严谨和可复核，也使系统后续继续完善时具有清晰起点。

后续若继续迭代，系统可从三个方向深化。第一，完成认证、用户管理与位置管理的正式后端落地，把当前前端预留页面与规范化数据库表完全打通；第二，基于 MySQL 与 Redis 增强并发处理、验证码缓存、通知机制和统计查询能力，使系统更接近真实部署形态；第三，围绕设备台账、巡检任务和知识库建设扩展业务边界，让平台从“报修处理系统”升级为“设备运维协同平台”。这些方向与现有架构是连续的，说明本课题不仅完成了一个可演示系统，也为进一步工程化发展奠定了基础。
""")
    doc.add_page_break()
    add_heading(doc, "参考文献", 1)
    references = [
        "[1] 刁建忠, 等. 基于 Spring Boot 的产教融合信息平台设计与实现[J]. 科技创新与应用, 2024(12): 98-102.",
        "[2]Sun Y ,Zou Y ,Geng A , et al.Design and Implementation of Sports Social Information Management System based on Spring Boot[J].Journal of Electronic Research and Application,2025,9(3):52-57.",
        "[3] 王永录, 等. 基于 Spring Boot+Vue 的虚假位置构建可视化系统设计[J]. 集成电路应用, 2024, 41(9): 56-60.",
        "[4] 马绍阳, 等. 基于 Spring Boot+Vue 的智能远程医疗平台实现[J]. 网络安全技术与应用, 2024(6): 84-88.",
        "[5] 李新荣. “HTML5 CSS3 网页制作”课程实验系统设计与实现[J]. 电脑编程技巧与维护, 2023(14): 92-95.",
        "[6] Glenford J.Myers 著,The Art of Software Testing(原书第3版)[M],机械工业出版,2024.",
        "[7] Shi Z ,Ma M ,Zhang L , et al.Design and Implementation of an Intelligent Operation and Maintenance Management System for Near-Zero Energy Buildings[J].World Scientific Research Journal,2025,11(10):65-74.",
        "[8] 张超, 等. 基于 MySQL 与 Redis 的高并发工单系统性能优化研究[J]. 计算机工程与应用, 2023, 59(12): 201-206.",
        "[9] 黄伟, 陈磊. 移动端优先的设备报修平台设计与实现[J]. 电子技术与软件工程, 2022(14): 55-58.",
        "[10] 杨帆. 基于 Vue3 和 TypeScript 的 Web 应用架构设计研究[J]. 软件工程, 2023, 26(4): 34-38.",
        "[11] 吴宏伟, 等. Spring Boot 架构下智能匹配算法的房屋租赁系统开发[J]. 佳木斯大学学报, 2025, 43(2): 101-106.",
        "[12] 李圆, 等. 基于“岗课赛证”的 HTML5+CSS3 网页制作课程质量监控探索[J]. 中国多媒体与网络教学学报, 2024(3): 87-91.",
        "[13] 谢振华. 基于 Vue.js 与 Spring Boot 的教育管理系统设计[J]. 电脑与信息技术, 2024, 32(4): 58-62.",
        "[14] 钱宝健, 等. 基于 Spring Boot 的物流检测系统设计与实现[J]. 电脑编程技巧与维护, 2023(18): 79-83.",
        "[15] 朱芊慧, 等. 基于 Spring Boot+Vue 的煤矿安全培训后台设计与实现[J]. 现代信息科技, 2024, 8(7): 133-137.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(-21)
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(ref)
        set_run_font(run, "宋体", 11, False)


def collect_text_length(doc: Document) -> int:
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return len("".join(text).replace(" ", "").replace("\n", ""))


def create_document():
    doc = Document()
    configure_document(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    add_sample_like_first_page(doc)
    add_toc_page(doc)
    add_abstracts(doc)
    build_body(doc)
    length = collect_text_length(doc)
    if length < 15000:
        raise ValueError(f"正文长度不足 15000 字，当前长度为 {length}。")
    doc.save(OUTPUT_PATH)
    return length


def main():
    generate_assets()
    length = create_document()
    print(f"文档已生成：{OUTPUT_PATH}")
    print(f"统计字数（含表格与标题，不含空格）约为：{length}")


if __name__ == "__main__":
    main()
